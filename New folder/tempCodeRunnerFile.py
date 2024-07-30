from docx import Document
from docx.shared import Pt

# Create a new Document
doc = Document()

# Define styles
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Add header
header = doc.add_heading(level=1)
header_run = header.add_run("Prabhat Zade\nData Scientist/Data Engineer\n")
header_run.bold = True

# Contact Information
contact_info = "+91 7020882015\nprabhatzade@gmail.com\nHyderabad, Telangana\nLinkedIn\n"
doc.add_paragraph(contact_info)

# Profile Summary
doc.add_heading("Profile Summary", level=2)
profile_summary = (
    "Results-oriented Data Scientist with expertise in Python, SQL, machine learning, and Hadoop. "
    "Proven track record of driving data-driven decisions and innovation. Key contributions to projects like FaceNet "
    "and Propensity Modeling demonstrate proficiency in face recognition and predictive analytics. Recognized for "
    "exceptional problem-solving skills and pioneering technology implementations. Delivered impactful solutions that "
    "enhanced user experiences and drove project success. Seeking to leverage my skills and experience in a challenging "
    "Data Science role, building on a strong foundation in Python, SQL, and technology implementations."
)
doc.add_paragraph(profile_summary)

# Skills
doc.add_heading("Skills", level=2)
skills = (
    "- **Programming:** Python, SQL\n"
    "- **Machine Learning:** Predictive Modeling, Deep Learning, CNN\n"
    "- **Data Processing:** Hadoop, Hive, HBase, Spark\n"
    "- **NLP:** Natural Language Processing, Langchain\n"
    "- **Computer Vision:** OpenCV\n"
    "- **Visualization:** Tableau\n"
    "- **Statistics:** Statistical Analysis\n"
    "- **Scripting:** Unix Shell Scripting\n"
    "- **Version Control:** GitHub\n"
    "- **Soft Skills:** Problem-Solving, Team Collaboration, Communication\n"
)
doc.add_paragraph(skills)

# Work Experience
doc.add_heading("Work Experience", level=2)

# Project Engineer
job_title1 = "Project Engineer | Centre for Development of Advanced Computing (C-DAC R&D), Hyderabad | Dec '22 – Present"
doc.add_heading(job_title1, level=3)
contributions1 = (
    "- Led projects such as FaceNet and Propensity Modeling, demonstrating expertise in face recognition, predictive analytics, NLP, and the Langchain Framework.\n"
    "- Drove data-driven decisions through innovative solutions, recognized for implementing advanced machine learning algorithms.\n"
    "- Excelled in data preprocessing, modeling, and evaluation, ensuring accuracy and reliability of analytical insights.\n"
    "- Developed and maintained data pipelines for large-scale data processing, optimizing efficiency and scalability.\n"
    "- Proficient in deploying and managing complex data infrastructure using Hadoop and SQL for robust data management.\n"
    "- Utilized tools like Excel for in-depth data exploration and visualization.\n"
    "- Engaged in continuous learning to stay updated with the latest advancements in data science and engineering.\n"
    "- Applied problem-solving skills to manage multiple projects concurrently, ensuring timely and high-quality delivery.\n"
    "- Collaborated with cross-functional teams to translate business requirements into technical solutions, fostering a culture of innovation and excellence.\n"
)
doc.add_paragraph(contributions1)

# Associate Data Engineer Intern
job_title2 = "Associate Data Engineer Intern | phData Solutions Pvt. Ltd. | Jun 2022 - Jul 2022"
doc.add_heading(job_title2, level=3)
contributions2 = (
    "- Completed a comprehensive internship focusing on Python, SQL, and Snowflake.\n"
    "- Designed, developed, tested, and deployed data pipelines, gaining hands-on experience in end-to-end data engineering.\n"
    "- Applied Agile methodologies and JIRA Scrum frameworks for effective project management and collaboration.\n"
    "- Contributed to enhancing data quality and efficiency by adopting best practices for data management and optimization.\n"
    "- Participated actively in team meetings, leveraging collaborative problem-solving skills to address technical challenges.\n"
    "- Demonstrated adaptability and a quick learning ability, assimilating new concepts and technologies to support project objectives.\n"
)
doc.add_paragraph(contributions2)

# Education
doc.add_heading("Education", level=2)
education = (
    "**P.G. Diploma in Big Data Analytics (PG-DBDA) | Centre for Development of Advanced Computing (C-DAC), Bangalore | Sept '21 - Apr '22 | Score: 67%**\n\n"
    "**Bachelors in Engineering (Computer Science and Engineering) | Sant Gadge Baba Amravati University, Amravati, Maharashtra | Aug '16 - Aug '21 | Score: 92.74%**\n"
)
doc.add_paragraph(education)

# Interests
doc.add_heading("Interests", level=2)
interests = "- New Technology Exploration\n- Gym\n- Traveling"
doc.add_paragraph(interests)

# Languages
doc.add_heading("Languages", level=2)
languages = "- English\n- Hindi\n- Marathi"
doc.add_paragraph(languages)

# Save the document
file_path = "/mnt/data/Prabhat_Zade_Resume.docx"
doc.save(file_path)

file_path
